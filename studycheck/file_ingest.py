from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path

@dataclass(frozen=True)
class StudyDocument:
    name: str
    text: str
    size_bytes: int
    media_type: str

class IngestError(ValueError): pass

_ALLOWED={'.txt':'text/plain','.pdf':'application/pdf','.docx':'application/vnd.openxmlformats-officedocument.wordprocessingml.document'}

def _extract(path:Path,suffix:str)->str:
    if suffix=='.txt':
        return path.read_text(encoding='utf-8')
    try:
        if suffix=='.pdf':
            from pypdf import PdfReader
            return '\n'.join(page.extract_text() or '' for page in PdfReader(str(path)).pages)
        if suffix=='.docx':
            from docx import Document
            doc=Document(str(path))
            return '\n'.join(p.text for p in doc.paragraphs)
    except ImportError as exc:
        raise IngestError(f'{suffix} parser dependency is not installed') from exc
    except Exception as exc:
        raise IngestError(f'failed to parse {suffix} document') from exc
    raise IngestError('unsupported file type')

def ingest_text(path:str|Path,max_bytes:int=5_000_000)->StudyDocument:
    p=Path(path)
    if max_bytes<=0: raise ValueError('max_bytes must be positive')
    if not p.is_file(): raise IngestError('input file not found')
    suffix=p.suffix.lower()
    if suffix not in _ALLOWED: raise IngestError('unsupported file type')
    size=p.stat().st_size
    if size>max_bytes: raise IngestError('input file exceeds size limit')
    try: text=_extract(p,suffix)
    except UnicodeDecodeError as exc: raise IngestError('input must be UTF-8 text') from exc
    if not text.strip(): raise IngestError('input document is empty')
    return StudyDocument(p.name,text,size,_ALLOWED[suffix])
